from pdfminer.high_level import extract_pages
from pdfminer.layout import LTTextContainer
from pdfminer.layout import LAParams
from docx import Document
import datetime
from datetime import date
import os
import os.path
import Cocoa

# ================= MAIN ====================
def main():
    now = datetime.datetime.now()
    date = now.strftime("%d-%m-%Y %H%p")

    # base_path = os.getcwd()
    base_path = '/'.join(str(Cocoa.NSBundle.mainBundle().bundlePath()).split('/')[:-1]) # Mac version
    folder_path = os.path.join(base_path, 'invoices')

    PDF_paths = []
    while not PDF_paths:
        for dirpath, dirnames, filenames in os.walk(folder_path, onerror=walk_error_handler):
            for pdf in filenames:
                path = os.path.join(dirpath, pdf)
                if path.endswith('.pdf'):
                    PDF_paths.append(path)

        if not PDF_paths:
            base_path = os.path.join(os.environ.get('HOME'), 'OneDrive', 'Desktop') # USERPROFILE
            folder_path = os.path.join(base_path, 'invoices')

    txt_path = os.path.join(base_path, date + '.txt')
    tally_path = os.path.join(base_path, date + ' Order count.txt')
    doc_path = os.path.join(base_path, date + '.docx')

    PDF_paths_ordered = order_pdf_paths(PDF_paths)
    tally_dict = {'FT': 0, 'SWISS': 0, 'MULTITOOL': 0, 'LEDLENSER': 0, 'GERBER': 0, 'BUCK': 0, 'TOTAL': 0}

    for i, paths in enumerate(PDF_paths_ordered):
        if paths:
            pages = read_pdf(paths)
            text = []

            if i == 0:
                order_tally = 0
                for page in pages:
                    words = get_words_swiss(page)

                    for item in page:
                        if 'Order number' in item:
                            order_tally += 1
                            break

                    if not words:
                        continue

                    for item in words:
                        text.append(item)

                text = sort_engravings_FT(text)
                FT_titles = ('F&T - WOOD\n', 'F&T - METAL\n', 'F&T - GLASS\n')
                write_to_doc(text, FT_titles, txt_path, doc_path)

                tally_dict['FT'] += order_tally

            elif i == 1:
                order_tally = 0
                for page in pages:
                    words = get_words_swiss(page)

                    for item in page:
                        if 'Order number' in item:
                            order_tally += 1
                            break

                    if not words:
                        continue

                    for item in words:
                        text.append(item)

                text = sort_engravings_SWI(text)
                swiss_titles = ('SWISS - METAL\n', 'SWISS - PLASTIC\n')
                write_to_doc(text, swiss_titles, txt_path, doc_path)

                tally_dict['SWISS'] += order_tally

            else:
                order_tally = 0
                for page in pages:
                    words = get_words_txt(page)

                    for item in page:
                        if 'Order number' in item:
                            order_tally += 1
                            break

                    if not words:
                        continue

                    text.append(words)

                if i == 2:
                    MULTI_title = 'MULTITOOL\n'
                    write_to_doc(text, MULTI_title, txt_path, doc_path)
                    tally_dict['MULTITOOL'] += order_tally

                if i == 3:
                    LED_title = 'LEDLENSER\n'
                    write_to_doc(text, LED_title, txt_path, doc_path)
                    tally_dict['LEDLENSER'] += order_tally

                if i == 4:
                    GERBER_title = 'GERBER\n'
                    write_to_doc(text, GERBER_title, txt_path, doc_path)
                    tally_dict['GERBER'] += order_tally

                if i == 5:
                    BUCK_title = 'BUCK\n'
                    write_to_doc(text, BUCK_title, txt_path, doc_path)
                    tally_dict['BUCK'] += order_tally
        else:
            continue

    write_tally(tally_dict, tally_path)


def read_pdf_page(inPDF):
    text = []

    for page_layout in extract_pages(inPDF):
        if len(page_layout) > 30:
            for element in page_layout:
                if isinstance(element, LTTextContainer):
                    text.append(element.get_text())
            return text
        # Redundant?
        else:
            continue


def read_pdf(paths):
    pages = []
    text = []
    laparams = LAParams(boxes_flow=1)

    for pdf in paths:
        for page_layout in extract_pages(pdf, laparams=laparams):
            for element in page_layout:
                if isinstance(element, LTTextContainer):
                    text.append(element.get_text())

            pages.append(text)
            text = []

    return pages


def get_words_txt(page):
    words = [i for i in page if 'Engraving Message' in i or ('Text' in i and 'Handwritten' not in i)]
    if not words:
        return None


    for i, item in enumerate(words):
        if 'See Email' in item:
            return None

        start_idx = item.find('Text')
        if start_idx < 0:
            start_idx = item.find('Message:') + 9
        else:
            start_idx += 7

        item = item[start_idx:]
        # print('=======================')
        # print(item)
        if '||' in item:
            end_idx = item.rfind('\n', item.find(':'), find_2nd(item, ':'))
        else:
            end_idx = item.rfind('\n', 0, item.find(':'))

        item = item[:end_idx].rstrip()
        if '\n' in item:
            item = item.replace('\n', ' ')
        item = item + '\n'
        words[i] = item

    if len(words) == 1:
        words = words[0]

    return words


def get_words_swiss(page):
    word_list = []
    engravings = []

    words = [i for i in page if 'Engraving Message' in i]
    if not words:
        return None
    # print('=======================')
    # print(words)

    for i, item in enumerate(words):
        if 'See Email' in item:
            return None

        start_idx = item.find('Text')
        if start_idx == -1:
            start_idx = item.find('Message:') + 9
        else:
            start_idx += 7

        # text = item[start_idx : item.find('Engraving Font')].rstrip()
        text = item[start_idx:]

        # end_idx = text.rfind('\n', 0, text.find(':'))
        if '||' in text:
            end_idx = text.rfind('\n', text.find(':'), find_2nd(text, ':'))
        else:
            end_idx = text.rfind('\n', 0, text.find(':'))

        text = text[:end_idx].rstrip()

        if '\n' in text:
            text = text.replace('\n', ' ')
        text = text + '\n'

        position = item[(item.find('options')) + 9:item.find('Engraving Message')]
        engravings = [position, text]
        word_list.append(engravings)

    return word_list

def find_2nd(string, substring):
   return string.find(substring, string.find(substring) + 1)


# def get_words_FT(page):
#     word_list = []
#     engravings = []
#
#     words = [i for i in page if 'Engraving Message' in i]
#     if not words:
#         return None
#
#     for i, item in enumerate(words):
#         start_idx = []
#         print('\n', item)
#         for m in re.finditer('Text', item):
#             print('Text found', m.start(), m.end())
#             start_idx.append(m.start() + 7)
#
#         if not start_idx:
#             for m in re.finditer('Message:', item):
#                 print('Message found', m.start(), m.end())
#                 start_idx.append(m.start() + 9)
#
#         for idx in start_idx:
#             text = item[idx : item.find('Engraving Font')].rstrip()
#             if '\n' in text:
#                 text = text.replace('\n', ' ')
#             text = text + '\n'
#             position = item[(item.find('options')) + 9:item.find('Engraving Message')]
#             engravings = [position, text]
#             word_list.append(engravings)
#
#     return word_list


def get_store(page):
    store_temp = None

    for i in page:
        if 'Order Reference' in i or 'Order number' in i:
            store_temp = i
            if 'Order number' in store_temp:
                store_temp = store_temp[store_temp.find('Order number'):]

            idx = (store_temp.find(':') + 1)
            if store_temp[idx] == ' ':
                store_temp = store_temp[idx + 1:]
            else:
                store_temp = store_temp[idx:]
            break

    return store_temp


def sort_engravings_FT(text):
    check_metal = ('Blade', 'Knife Blade', 'Blades', 'Knife Blades', 'Brass', 'Silver', 'Plate', 'Hip', 'Watch', 'Front', 'Flask', 'Body', 'Side', 'Back', 'Cufflink', 'Base', 'Mist Sprayer', 'Top', 'Cup', 'Nail')
    wood = []
    metal = []
    glass = []

    for items in text:
        p = items[0]
        e = items[1]

        if any(i in p for i in check_metal) and 'Glass' not in p and 'Wooden' not in p and 'Pewter' not in p and 'Leather' not in p and 'Upper' not in p and 'Wallet' not in p and 'Board' not in p:
            metal.append(e)
        elif 'Glass' in p or 'Flute' in p or 'Pewter' in p:
            glass.append(e)
        else:
            wood.append(e)

    return(wood, metal, glass)


def sort_engravings_SWI(text):
    plastic = []
    metal = []

    for items in text:
        w = items[0]
        e = items[1]

        if 'Handle' in w or 'Leather' in w or 'Pouch' in w:
            plastic.append(e)
        else:
            metal.append(e)

    return (metal, plastic)


def write_to_files_FTSWI(text, titles, doc, txt_file):
    p = doc.add_paragraph()

    for i, t in enumerate(titles):
        p.add_run(t).bold = True
        txt_file.write('\n' + t)

        for item in text[i]:
            p.add_run(item)
            txt_file.write(item)


def write_to_files(text, title, doc, txt_file):
    p = doc.add_paragraph()

    p.add_run(title).bold = True
    txt_file.write('\n' + title)
    for item in text:
        if isinstance(item, list):
            for w in item:
                p.add_run(w)
                txt_file.write(w)
        else:
            p.add_run(item)
            txt_file.write(item)


def write_to_doc(text, titles, txt_path, doc_path):
    try:
        with open(doc_path, 'rb') as docx_f:
            document = Document(docx_f)
    except:
        document = Document()

    txt_fa = open(txt_path, 'a', encoding='utf-8')

    if 'F&T' in titles[0] or 'SWISS' in titles[0]:
        write_to_files_FTSWI(text, titles, document, txt_fa)
    else:
        write_to_files(text, titles, document, txt_fa)

    txt_fa.close()
    document.save(doc_path)


def write_tally(dict, tally_path):
    text_file = open(tally_path, 'w', encoding='utf-8')
    dict['TOTAL'] = sum(dict.values())

    for key, value in dict.items():
        text_file.write(key + ' - ' + str(value) + '\n')

    text_file.close()


def order_pdf_paths(PDF_paths):
    store = ''
    FT = []
    SWISS = []
    MULTI = []
    LED = []
    GERBER = []
    BUCK = []

    for i, pdf in enumerate(PDF_paths):
        page = read_pdf_page(pdf)
        store = get_store(page)

        if store.startswith('00'):
            FT.append(PDF_paths[i])

        elif store.startswith('10'):
            MULTI.append(PDF_paths[i])

        elif store.startswith('13'):
            LED.append(PDF_paths[i])

        elif store.startswith('18'):
            SWISS.append(PDF_paths[i])

        elif store.startswith('6'):
            GERBER.append(PDF_paths[i])

        elif store.startswith('2'):
            BUCK.append(PDF_paths[i])

    list2d = (tuple(FT), tuple(SWISS), tuple(MULTI), tuple(LED), tuple(GERBER), tuple(BUCK))

    return list2d


def walk_error_handler(exception_instance):
    print(exception_instance)
    if 'OneDrive' in folder_path:
        quit()


main()
